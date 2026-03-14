"""
WordPress Publisher for Overlappr
Publishes blog post drafts to WordPress via REST API
Supports image extraction from PDF and DOCX files
"""

import json
import os
import tempfile
import requests
from requests.auth import HTTPBasicAuth

# Optional imports for document processing
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    from bs4 import BeautifulSoup
    HAS_BS4 = True
except ImportError:
    HAS_BS4 = False

def load_config():
    """Load WordPress credentials from config file"""
    with open('config.json', 'r') as f:
        return json.load(f)


def upload_image_to_wordpress(image_path, alt_text=""):
    """
    Upload an image to WordPress Media Library

    Args:
        image_path: Path to the image file
        alt_text: Alternative text for the image

    Returns:
        dict with image URL and ID if successful, None if failed
    """
    config = load_config()
    url = f"{config['wordpress_url']}/wp-json/wp/v2/media"

    filename = os.path.basename(image_path)

    # Determine content type based on extension
    ext = os.path.splitext(filename)[1].lower()
    content_types = {
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.gif': 'image/gif',
        '.webp': 'image/webp'
    }
    content_type = content_types.get(ext, 'image/jpeg')

    try:
        with open(image_path, 'rb') as img_file:
            headers = {
                'Content-Disposition': f'attachment; filename="{filename}"',
                'Content-Type': content_type
            }

            response = requests.post(
                url,
                headers=headers,
                data=img_file.read(),
                auth=HTTPBasicAuth(config['username'], config['application_password']),
                timeout=60
            )

            if response.status_code == 201:
                media = response.json()
                print(f"   ✅ Uploaded: {filename}")
                return {
                    'id': media['id'],
                    'url': media['source_url'],
                    'filename': filename
                }
            else:
                print(f"   ❌ Failed to upload {filename}: {response.status_code}")
                return None

    except Exception as e:
        print(f"   ❌ Error uploading {filename}: {e}")
        return None


def extract_images_from_pdf(pdf_path):
    """
    Extract all images from a PDF file

    Args:
        pdf_path: Path to the PDF file

    Returns:
        List of paths to extracted image files
    """
    if not HAS_PYMUPDF:
        print("   ⚠️ PyMuPDF not installed. Cannot extract images from PDF.")
        return []

    image_paths = []
    temp_dir = tempfile.mkdtemp(prefix="wp_images_")

    try:
        doc = fitz.open(pdf_path)

        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images(full=True)

            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]

                image_filename = f"image_p{page_num + 1}_{img_index + 1}.{image_ext}"
                image_path = os.path.join(temp_dir, image_filename)

                with open(image_path, 'wb') as img_file:
                    img_file.write(image_bytes)

                image_paths.append(image_path)

        doc.close()
        print(f"   📷 Extracted {len(image_paths)} images from PDF")

    except Exception as e:
        print(f"   ❌ Error extracting images from PDF: {e}")

    return image_paths


def extract_images_from_docx(docx_path):
    """
    Extract all images from a DOCX file

    Args:
        docx_path: Path to the DOCX file

    Returns:
        List of tuples (image_path, relationship_id)
    """
    if not HAS_DOCX:
        print("   ⚠️ python-docx not installed. Cannot extract images from DOCX.")
        return []

    image_paths = []
    temp_dir = tempfile.mkdtemp(prefix="wp_images_")

    try:
        doc = Document(docx_path)

        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                image_data = rel.target_part.blob
                image_ext = os.path.splitext(rel.target_ref)[1]

                image_filename = f"image_{rel_id}{image_ext}"
                image_path = os.path.join(temp_dir, image_filename)

                with open(image_path, 'wb') as img_file:
                    img_file.write(image_data)

                image_paths.append((image_path, rel_id))

        print(f"   📷 Extracted {len(image_paths)} images from DOCX")

    except Exception as e:
        print(f"   ❌ Error extracting images from DOCX: {e}")

    return image_paths


def convert_docx_to_html(docx_path, uploaded_images=None):
    """
    Convert a DOCX file to HTML with all formatting preserved

    Args:
        docx_path: Path to the DOCX file
        uploaded_images: Dict mapping rel_id to WordPress image URLs

    Returns:
        dict with 'title', 'html_content', and 'excerpt'
    """
    if not HAS_DOCX:
        print("   ⚠️ python-docx not installed. Cannot convert DOCX.")
        return None

    if uploaded_images is None:
        uploaded_images = {}

    try:
        doc = Document(docx_path)
        html_parts = []
        title = ""
        first_paragraph = ""

        # Build hyperlink map from relationships
        hyperlink_rels = {}
        for rel_id, rel in doc.part.rels.items():
            if "hyperlink" in rel.reltype:
                hyperlink_rels[rel_id] = rel.target_ref

        for para in doc.paragraphs:
            # Determine heading level
            style_name = para.style.name.lower() if para.style else ""

            # Extract paragraph content with formatting
            para_html = process_paragraph_runs(para, hyperlink_rels, uploaded_images)

            if not para_html.strip():
                continue

            # Handle headings
            if "heading 1" in style_name or style_name == "title":
                if not title:
                    title = para.text.strip()
                html_parts.append(f"<h1>{para_html}</h1>")
            elif "heading 2" in style_name:
                html_parts.append(f"<h2>{para_html}</h2>")
            elif "heading 3" in style_name:
                html_parts.append(f"<h3>{para_html}</h3>")
            elif "heading 4" in style_name:
                html_parts.append(f"<h4>{para_html}</h4>")
            elif "heading 5" in style_name:
                html_parts.append(f"<h5>{para_html}</h5>")
            elif "heading 6" in style_name:
                html_parts.append(f"<h6>{para_html}</h6>")
            else:
                # Regular paragraph
                if not first_paragraph and para.text.strip():
                    first_paragraph = para.text.strip()[:200]
                html_parts.append(f"<p>{para_html}</p>")

        # Process tables
        for table in doc.tables:
            table_html = convert_table_to_html(table, hyperlink_rels)
            html_parts.append(table_html)

        html_content = "\n".join(html_parts)

        print(f"   📄 Converted DOCX to HTML")
        print(f"   📝 Title: {title[:50]}..." if len(title) > 50 else f"   📝 Title: {title}")

        return {
            'title': title,
            'html_content': html_content,
            'excerpt': first_paragraph
        }

    except Exception as e:
        print(f"   ❌ Error converting DOCX: {e}")
        import traceback
        traceback.print_exc()
        return None


def process_paragraph_runs(paragraph, hyperlink_rels, uploaded_images=None):
    """
    Process paragraph runs to extract text with formatting and hyperlinks

    Args:
        paragraph: A python-docx paragraph object
        hyperlink_rels: Dict mapping relationship IDs to URLs
        uploaded_images: Dict mapping rel_id to WordPress image URLs

    Returns:
        HTML string with formatting preserved
    """
    from docx.oxml.ns import qn

    if uploaded_images is None:
        uploaded_images = {}

    html_parts = []

    for child in paragraph._element:
        # Handle hyperlinks
        if child.tag == qn('w:hyperlink'):
            rel_id = child.get(qn('r:id'))
            url = hyperlink_rels.get(rel_id, "#")

            # Get hyperlink text
            link_text = ""
            for run in child.findall(qn('w:r')):
                text_elem = run.find(qn('w:t'))
                if text_elem is not None and text_elem.text:
                    link_text += text_elem.text

            if link_text:
                html_parts.append(f'<a href="{url}">{link_text}</a>')

        # Handle regular runs
        elif child.tag == qn('w:r'):
            run_text = ""
            for elem in child:
                if elem.tag == qn('w:t'):
                    if elem.text:
                        run_text += elem.text
                # Handle images in runs
                elif elem.tag == qn('w:drawing'):
                    # Try to find the image relationship
                    blip = elem.find('.//' + qn('a:blip'))
                    if blip is not None:
                        embed_id = blip.get(qn('r:embed'))
                        if embed_id and embed_id in uploaded_images:
                            img_url = uploaded_images[embed_id]
                            html_parts.append(f'<img src="{img_url}" alt="Image" />')

            if run_text:
                # Check for formatting
                rPr = child.find(qn('w:rPr'))
                if rPr is not None:
                    is_bold = rPr.find(qn('w:b')) is not None
                    is_italic = rPr.find(qn('w:i')) is not None
                    is_underline = rPr.find(qn('w:u')) is not None

                    if is_bold:
                        run_text = f"<strong>{run_text}</strong>"
                    if is_italic:
                        run_text = f"<em>{run_text}</em>"
                    if is_underline:
                        run_text = f"<u>{run_text}</u>"

                html_parts.append(run_text)

    return "".join(html_parts)


def convert_table_to_html(table, hyperlink_rels=None):
    """
    Convert a DOCX table to HTML

    Args:
        table: A python-docx Table object
        hyperlink_rels: Dict mapping relationship IDs to URLs

    Returns:
        HTML table string
    """
    if hyperlink_rels is None:
        hyperlink_rels = {}

    html = ["<table>"]

    for row_idx, row in enumerate(table.rows):
        html.append("  <tr>")
        for cell in row.cells:
            tag = "th" if row_idx == 0 else "td"
            cell_text = cell.text.strip()
            html.append(f"    <{tag}>{cell_text}</{tag}>")
        html.append("  </tr>")

    html.append("</table>")
    return "\n".join(html)


def convert_pdf_to_html(pdf_path, uploaded_images=None):
    """
    Convert a PDF file to HTML with formatting preserved

    Args:
        pdf_path: Path to the PDF file
        uploaded_images: List of uploaded image info

    Returns:
        dict with 'title', 'html_content', and 'excerpt'
    """
    if not HAS_PYMUPDF:
        print("   ⚠️ PyMuPDF not installed. Cannot convert PDF.")
        return None

    if uploaded_images is None:
        uploaded_images = []

    try:
        doc = fitz.open(pdf_path)
        html_parts = []
        title = ""
        first_paragraph = ""
        all_links = []

        # First pass: collect all links
        for page_num in range(len(doc)):
            page = doc[page_num]
            links = page.get_links()
            for link in links:
                if link.get('uri'):
                    all_links.append({
                        'page': page_num,
                        'rect': link['from'],
                        'uri': link['uri']
                    })

        # Second pass: extract text with structure
        for page_num in range(len(doc)):
            page = doc[page_num]

            # Get text blocks with position info
            blocks = page.get_text("dict")["blocks"]

            for block in blocks:
                if block['type'] == 0:  # Text block
                    for line in block.get('lines', []):
                        line_text = ""
                        line_size = 0
                        is_bold = False

                        for span in line.get('spans', []):
                            span_text = span.get('text', '').strip()
                            if span_text:
                                line_text += span_text + " "
                                line_size = max(line_size, span.get('size', 12))
                                if 'bold' in span.get('font', '').lower():
                                    is_bold = True

                        line_text = line_text.strip()
                        if not line_text:
                            continue

                        # Check if this text has a link
                        line_rect = fitz.Rect(line['bbox'])
                        for link in all_links:
                            if link['page'] == page_num:
                                link_rect = fitz.Rect(link['rect'])
                                if line_rect.intersects(link_rect):
                                    line_text = f'<a href="{link["uri"]}">{line_text}</a>'
                                    break

                        # Determine if heading based on font size
                        if line_size >= 24:
                            if not title:
                                title = line_text.replace('<a href=', '').split('>')[1].split('<')[0] if '<a href=' in line_text else line_text
                            html_parts.append(f"<h1>{line_text}</h1>")
                        elif line_size >= 18:
                            html_parts.append(f"<h2>{line_text}</h2>")
                        elif line_size >= 14 and is_bold:
                            html_parts.append(f"<h3>{line_text}</h3>")
                        else:
                            if not first_paragraph:
                                first_paragraph = line_text[:200]
                            html_parts.append(f"<p>{line_text}</p>")

        doc.close()

        # Add uploaded images at the end (or user can reposition in WordPress)
        if uploaded_images:
            html_parts.append("<h2>Images</h2>")
            for img in uploaded_images:
                html_parts.append(f'<img src="{img["url"]}" alt="{img["filename"]}" />')

        html_content = "\n".join(html_parts)

        print(f"   📄 Converted PDF to HTML")
        print(f"   📝 Title: {title[:50]}..." if len(title) > 50 else f"   📝 Title: {title}")
        print(f"   🔗 Found {len(all_links)} hyperlinks")

        return {
            'title': title,
            'html_content': html_content,
            'excerpt': first_paragraph
        }

    except Exception as e:
        print(f"   ❌ Error converting PDF: {e}")
        import traceback
        traceback.print_exc()
        return None


def upload_images_to_wordpress(image_paths):
    """
    Upload multiple images to WordPress

    Args:
        image_paths: List of image file paths

    Returns:
        List of uploaded image info (url, id)
    """
    uploaded = []

    print(f"\n📤 Uploading {len(image_paths)} images to WordPress...")

    for path in image_paths:
        result = upload_image_to_wordpress(path)
        if result:
            uploaded.append(result)

    print(f"   ✅ Successfully uploaded {len(uploaded)}/{len(image_paths)} images")

    return uploaded


def process_document(file_path):
    """
    Process a PDF or DOCX file: extract images and upload to WordPress

    Args:
        file_path: Path to the PDF or DOCX file

    Returns:
        dict with uploaded_images list
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        image_paths = extract_images_from_pdf(file_path)
        # Convert to list format for compatibility
        image_paths = [(p, f"img_{i}") for i, p in enumerate(image_paths)]
    elif ext in ['.docx', '.doc']:
        image_paths = extract_images_from_docx(file_path)
    else:
        print(f"   ⚠️ Unsupported file type: {ext}")
        return {'uploaded_images': []}

    if image_paths:
        # Upload images and build mapping
        uploaded = []
        uploaded_map = {}
        print(f"\n📤 Uploading {len(image_paths)} images to WordPress...")
        for path, rel_id in image_paths:
            result = upload_image_to_wordpress(path)
            if result:
                uploaded.append(result)
                uploaded_map[rel_id] = result['url']
        print(f"   ✅ Successfully uploaded {len(uploaded)}/{len(image_paths)} images")
        return {'uploaded_images': uploaded, 'image_map': uploaded_map}

    return {'uploaded_images': [], 'image_map': {}}


def process_and_publish_document(file_path):
    """
    Full document processing: extract content, images, links, and publish to WordPress

    This is the main function for the /wpdraft command when a document is provided.

    Args:
        file_path: Path to the PDF or DOCX file

    Returns:
        dict with post info if successful, None if failed
    """
    print("\n" + "="*60)
    print("📄 PROCESSING DOCUMENT FOR WORDPRESS")
    print("="*60)
    print(f"\n📁 File: {os.path.basename(file_path)}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext not in ['.pdf', '.docx', '.doc']:
        print(f"   ❌ Unsupported file type: {ext}")
        print("   Supported formats: PDF, DOCX")
        return None

    # Step 1: Extract and upload images
    print("\n📷 Step 1: Extracting and uploading images...")
    doc_result = process_document(file_path)
    uploaded_images = doc_result.get('uploaded_images', [])
    image_map = doc_result.get('image_map', {})

    if uploaded_images:
        print(f"   ✅ {len(uploaded_images)} images uploaded to WordPress Media Library")
    else:
        print("   ℹ️ No images found in document")

    # Step 2: Convert document to HTML
    print("\n📝 Step 2: Converting document to HTML...")

    if ext == '.pdf':
        result = convert_pdf_to_html(file_path, uploaded_images)
    else:  # DOCX
        result = convert_docx_to_html(file_path, image_map)

    if not result:
        print("   ❌ Failed to convert document")
        return None

    title = result['title'] or "Untitled Post"
    html_content = result['html_content']
    excerpt = result['excerpt'] or ""

    print(f"   ✅ Document converted successfully")

    # Step 3: Publish to WordPress
    print("\n🚀 Step 3: Publishing to WordPress as draft...")

    post_result = create_draft_post(title, html_content, excerpt)

    if post_result:
        print("\n" + "="*60)
        print("🎉 SUCCESS! Document published to WordPress")
        print("="*60)
        print(f"\n   📝 Title: {title}")
        print(f"   🔗 Edit URL: {post_result['edit_url']}")
        print(f"   📷 Images uploaded: {len(uploaded_images)}")
        print(f"   📄 Status: Draft (ready for your review)")
        print("\n   Next steps:")
        print("   1. Open the edit URL above")
        print("   2. Review the content and formatting")
        print("   3. Add featured image if needed")
        print("   4. Optimize SEO settings")
        print("   5. Publish when ready!")
        print("="*60 + "\n")

        return {
            'post': post_result,
            'images_uploaded': len(uploaded_images),
            'title': title
        }

    return None


def test_connection():
    """Test connection to WordPress site"""
    config = load_config()

    # Test if we can reach the WordPress REST API
    url = f"{config['wordpress_url']}/wp-json/wp/v2/posts"

    try:
        response = requests.get(
            url,
            auth=HTTPBasicAuth(config['username'], config['application_password']),
            timeout=10
        )

        if response.status_code == 200:
            print("✅ Successfully connected to WordPress!")
            print(f"   Site: {config['wordpress_url']}")
            return True
        elif response.status_code == 401:
            print("❌ Authentication failed. Please check your username and application password.")
            return False
        else:
            print(f"❌ Connection failed with status code: {response.status_code}")
            return False

    except requests.exceptions.RequestException as e:
        print(f"❌ Connection error: {e}")
        return False

def create_draft_post(title, content, excerpt=""):
    """
    Create a draft post on WordPress

    Args:
        title: The blog post title
        content: The full HTML/text content of the post
        excerpt: Optional short description for SEO

    Returns:
        dict with post info if successful, None if failed
    """
    config = load_config()
    url = f"{config['wordpress_url']}/wp-json/wp/v2/posts"

    post_data = {
        'title': title,
        'content': content,
        'excerpt': excerpt,
        'status': 'draft'  # Always creates as draft for your review
    }

    try:
        response = requests.post(
            url,
            json=post_data,
            auth=HTTPBasicAuth(config['username'], config['application_password']),
            timeout=30
        )

        if response.status_code == 201:
            post = response.json()
            print(f"✅ Draft created successfully!")
            print(f"   Title: {post['title']['raw']}")
            print(f"   Post ID: {post['id']}")
            print(f"   Edit URL: {config['wordpress_url']}/wp-admin/post.php?post={post['id']}&action=edit")
            return {
                'id': post['id'],
                'title': post['title']['raw'],
                'edit_url': f"{config['wordpress_url']}/wp-admin/post.php?post={post['id']}&action=edit",
                'status': 'draft'
            }
        else:
            print(f"❌ Failed to create post: {response.status_code}")
            print(f"   Error: {response.text}")
            return None

    except requests.exceptions.RequestException as e:
        print(f"❌ Error creating post: {e}")
        return None

def publish_blog_post(title, content, excerpt=""):
    """
    Main function to publish a blog post as draft
    This is the function Claude Code will use to push your content
    """
    print("\n" + "="*50)
    print("📝 Publishing to WordPress...")
    print("="*50 + "\n")

    result = create_draft_post(title, content, excerpt)

    if result:
        print("\n" + "="*50)
        print("🎉 Your draft is ready for review!")
        print(f"   Go to: {result['edit_url']}")
        print("="*50 + "\n")

    return result

def extract_google_doc_id(url):
    """
    Extract document ID from a Google Docs URL.

    Supports formats:
      - https://docs.google.com/document/d/{ID}/edit?...
      - https://docs.google.com/document/d/{ID}/pub
      - https://docs.google.com/document/d/{ID}/view

    Returns:
        Document ID string, or None if not found
    """
    import re
    match = re.search(r'/document/d/([a-zA-Z0-9_-]+)', url)
    if match:
        return match.group(1)
    return None


def download_google_doc_html(doc_id):
    """
    Download a Google Doc as HTML using the public export URL.
    Works for any doc with link-sharing enabled (no auth required).

    Returns:
        HTML string, or None if failed
    """
    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=html"

    try:
        response = requests.get(export_url, timeout=30)

        if response.status_code == 200:
            print(f"   ✅ Downloaded Google Doc HTML ({len(response.content)} bytes)")
            return response.text
        elif response.status_code == 403:
            print("   ❌ Access denied. Make sure the Google Doc is shared publicly (Anyone with the link).")
            return None
        elif response.status_code == 404:
            print("   ❌ Document not found. Please check the URL.")
            return None
        else:
            print(f"   ❌ Failed to download document: HTTP {response.status_code}")
            return None

    except requests.exceptions.RequestException as e:
        print(f"   ❌ Connection error: {e}")
        return None


def process_google_doc_images(soup):
    """
    Find all images in parsed Google Doc HTML, download them,
    upload to WordPress Media Library, and replace src in-place.

    Args:
        soup: BeautifulSoup object of the document

    Returns:
        (updated soup, count of uploaded images)
    """
    import base64

    img_tags = soup.find_all('img')
    uploaded_count = 0
    temp_dir = tempfile.mkdtemp(prefix="gdoc_images_")

    print(f"   🖼️  Found {len(img_tags)} image(s) in document")

    for i, img in enumerate(img_tags):
        src = img.get('src', '')
        if not src:
            continue

        image_path = None

        try:
            # Case 1: base64 data URI  (data:image/png;base64,...)
            if src.startswith('data:image'):
                try:
                    header, encoded = src.split(',', 1)
                    ext = header.split('/')[1].split(';')[0]  # e.g. "png"
                    image_data = base64.b64decode(encoded)
                    image_path = os.path.join(temp_dir, f"gdoc_img_{i+1}.{ext}")
                    with open(image_path, 'wb') as f:
                        f.write(image_data)
                except Exception as e:
                    print(f"   ⚠️  Could not decode base64 image {i+1}: {e}")
                    continue

            # Case 2: external URL (Google CDN or other)
            elif src.startswith('http'):
                try:
                    r = requests.get(src, timeout=30)
                    if r.status_code == 200:
                        # Guess extension from Content-Type
                        content_type = r.headers.get('Content-Type', 'image/png')
                        ext_map = {
                            'image/jpeg': 'jpg', 'image/png': 'png',
                            'image/gif': 'gif', 'image/webp': 'webp'
                        }
                        ext = ext_map.get(content_type.split(';')[0].strip(), 'jpg')
                        image_path = os.path.join(temp_dir, f"gdoc_img_{i+1}.{ext}")
                        with open(image_path, 'wb') as f:
                            f.write(r.content)
                    else:
                        print(f"   ⚠️  Could not download image {i+1}: HTTP {r.status_code}")
                        continue
                except Exception as e:
                    print(f"   ⚠️  Error downloading image {i+1}: {e}")
                    continue

            else:
                continue

            # Upload to WordPress
            if image_path and os.path.exists(image_path):
                result = upload_image_to_wordpress(image_path)
                if result:
                    img['src'] = result['url']
                    img['alt'] = img.get('alt', '')
                    # Remove Google-specific style sizing constraints
                    if img.get('style'):
                        del img['style']
                    uploaded_count += 1

        except Exception as e:
            print(f"   ⚠️  Error processing image {i+1}: {e}")

    return soup, uploaded_count


def detect_and_convert_videos(soup):
    """
    Detect YouTube/Vimeo links wrapping images (how Google Docs exports embedded videos)
    and convert them to WordPress embed shortcodes.

    Returns:
        Updated soup
    """
    import re
    youtube_pattern = re.compile(
        r'(https?://)?(www\.)?(youtube\.com/watch\?v=|youtu\.be/)([a-zA-Z0-9_-]+)'
    )
    vimeo_pattern = re.compile(r'(https?://)?(www\.)?vimeo\.com/(\d+)')

    video_count = 0

    for a_tag in soup.find_all('a', href=True):
        href = a_tag['href']
        is_youtube = youtube_pattern.search(href)
        is_vimeo = vimeo_pattern.search(href)

        if (is_youtube or is_vimeo) and a_tag.find('img'):
            # Replace the linked thumbnail with a WordPress embed shortcode
            embed_tag = soup.new_tag('p')
            embed_tag.string = f'[embed]{href}[/embed]'
            a_tag.replace_with(embed_tag)
            video_count += 1

    if video_count:
        print(f"   🎬 Converted {video_count} video link(s) to WordPress embeds")

    return soup


def clean_google_doc_html(soup):
    """
    Strip Google-specific markup from exported HTML while keeping
    semantic structure: headings, paragraphs, bold, italic, underline,
    hyperlinks, tables, images.

    Returns:
        dict with 'title', 'html_content', 'excerpt'
    """
    # Remove style and script blocks
    for tag in soup.find_all(['style', 'script']):
        tag.decompose()

    # Strip all class attributes (Google internal CSS classes)
    for tag in soup.find_all(True):
        if tag.has_attr('class'):
            del tag['class']
        if tag.has_attr('id'):
            del tag['id']

    # Clean inline styles — keep only text-align, remove everything else
    for tag in soup.find_all(True):
        style = tag.get('style', '')
        if style:
            kept = []
            for rule in style.split(';'):
                rule = rule.strip()
                if rule.lower().startswith('text-align'):
                    kept.append(rule)
            if kept:
                tag['style'] = '; '.join(kept)
            else:
                del tag['style']

    # Extract title: prefer <title> tag, fall back to first h1
    title = ""
    title_tag = soup.find('title')
    if title_tag and title_tag.string:
        title = title_tag.string.strip()

    # Get body content
    body = soup.find('body')
    if not body:
        body = soup

    # If title came from <title> tag, also look for first h1 to use as post title
    first_h1 = body.find('h1')
    if first_h1 and first_h1.get_text(strip=True):
        title = first_h1.get_text(strip=True)

    # Extract excerpt from first non-empty paragraph
    excerpt = ""
    for p in body.find_all('p'):
        text = p.get_text(strip=True)
        if text:
            excerpt = text[:200]
            break

    # Render body inner HTML
    html_content = body.decode_contents().strip()

    return {
        'title': title,
        'html_content': html_content,
        'excerpt': excerpt
    }


def process_and_publish_google_doc(url):
    """
    Full pipeline for Mode 3: publicly shared Google Docs URL.

    Steps:
      1. Extract document ID from URL
      2. Download HTML export (no auth needed for public docs)
      3. Parse with BeautifulSoup
      4. Detect and convert embedded video links to WP embeds
      5. Download and upload all images to WordPress Media Library (in-place)
      6. Clean Google-specific HTML, preserve semantic structure
      7. Publish as WordPress draft

    Args:
        url: Publicly shared Google Docs URL

    Returns:
        dict with post info if successful, None if failed
    """
    if not HAS_BS4:
        print("❌ beautifulsoup4 is not installed.")
        print("   Run: pip3 install beautifulsoup4")
        return None

    print("\n" + "="*60)
    print("🌐 PROCESSING GOOGLE DOC FOR WORDPRESS")
    print("="*60)
    print(f"\n🔗 URL: {url}")

    # Step 1: Extract doc ID
    doc_id = extract_google_doc_id(url)
    if not doc_id:
        print("   ❌ Could not extract document ID from URL.")
        print("   Make sure it's a valid Google Docs link.")
        return None
    print(f"   📋 Document ID: {doc_id}")

    # Step 2: Download HTML
    print("\n📥 Step 1: Downloading Google Doc...")
    html = download_google_doc_html(doc_id)
    if not html:
        return None

    # Step 3: Parse
    soup = BeautifulSoup(html, 'html.parser')

    # Step 4: Detect and convert videos
    print("\n🎬 Step 2: Detecting embedded videos...")
    soup = detect_and_convert_videos(soup)

    # Step 5: Process and upload images
    print("\n📷 Step 3: Uploading images to WordPress...")
    soup, images_uploaded = process_google_doc_images(soup)
    if images_uploaded:
        print(f"   ✅ {images_uploaded} image(s) uploaded to WordPress Media Library")
    else:
        print("   ℹ️  No images found or uploaded")

    # Step 6: Clean HTML
    print("\n📝 Step 4: Cleaning and extracting content...")
    result = clean_google_doc_html(soup)
    title = result['title'] or "Untitled Post"
    html_content = result['html_content']
    excerpt = result['excerpt']

    print(f"   ✅ Title: {title[:60]}..." if len(title) > 60 else f"   ✅ Title: {title}")

    # Step 7: Publish to WordPress
    print("\n🚀 Step 5: Publishing to WordPress as draft...")
    post_result = create_draft_post(title, html_content, excerpt)

    if post_result:
        print("\n" + "="*60)
        print("🎉 SUCCESS! Google Doc published to WordPress")
        print("="*60)
        print(f"\n   📝 Title: {title}")
        print(f"   🔗 Edit URL: {post_result['edit_url']}")
        print(f"   📷 Images uploaded: {images_uploaded}")
        print(f"   📄 Status: Draft (ready for your review)")
        print("\n   Next steps:")
        print("   1. Open the edit URL above")
        print("   2. Review the content and formatting")
        print("   3. Add featured image if needed")
        print("   4. Optimize SEO settings")
        print("   5. Publish when ready!")
        print("="*60 + "\n")

        return {
            'post': post_result,
            'images_uploaded': images_uploaded,
            'title': title
        }

    return None


# Quick test when running directly
if __name__ == "__main__":
    print("\n🔌 Testing WordPress Connection...\n")
    test_connection()
